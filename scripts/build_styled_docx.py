"""
Publication-Grade Word Document Generator
Creates beautifully formatted Word documents with table callouts, cards, zebra tables, and souvenir matrices.
Robust against variable column counts, empty sections, and dynamic family/couple demographics.
"""
import os
import sys
import argparse
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, current_dir)
from parser import parse_travel_markdown

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b_elm = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val","single")}" w:sz="{edge_data.get("sz","4")}" w:space="0" w:color="{edge_data.get("color","auto")}"/>')
            tcBorders.append(b_elm)
        else:
            b_elm = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="none"/>')
            tcBorders.append(b_elm)
    tcPr.append(tcBorders)

def build_docx(input_path, output_path):
    data = parse_travel_markdown(input_path) if input_path and os.path.exists(input_path) else parse_travel_markdown("")
    doc = Document()
    
    # Page setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # 1. Hero Box
    t_hero = doc.add_table(rows=1, cols=1)
    t_hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_hero.autofit = False
    c_hero = t_hero.cell(0, 0)
    c_hero.width = Inches(6.8)
    set_cell_background(c_hero, "1E1B4B")
    set_cell_margins(c_hero, top=240, bottom=240, left=240, right=240)
    set_cell_border(c_hero)

    p = c_hero.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"🏰 {data.get('dates', '2026 行程规划')} · 精选定制度假")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(254, 205, 211)

    p2 = c_hero.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(data.get("title", "旅行度假保姆级规划"))
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(255, 255, 255)

    p3 = c_hero.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(data.get("subtitle", "轻松慢游 · 舒适休整 · 老少/双人惬意探索"))
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_callout(text_runs, bg_color="FFF1F2", border_color="E11D48"):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        c = t.cell(0, 0)
        c.width = Inches(6.8)
        set_cell_background(c, bg_color)
        set_cell_margins(c, top=120, bottom=120, left=160, right=160)
        set_cell_border(c, left=dict(sz=16, color=border_color))
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for txt, is_b, sz, col in text_runs:
            run = p.add_run(txt)
            run.font.bold = is_b
            run.font.size = Pt(sz)
            if col: run.font.color.rgb = col
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_heading(title, icon="📌"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{icon} {title}")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)

    # Core principle callout
    add_callout([
        ("💡 核心出行原则：", True, 10, RGBColor(225, 29, 72)),
        ("不赶路、不特种兵，合理规划节奏与休整，兼顾品质住宿、地道老饕美食与高品质在地文化体验！", False, 9.5, RGBColor(71, 85, 105))
    ])

    # 2. Outfits
    outfits = data.get("outfits", [])
    if outfits:
        add_heading(f"出行 {len(outfits)} 人穿搭卡片（色彩美学 · 拍照高级）", "👗")
        t_outfit = doc.add_table(rows=1, cols=len(outfits))
        t_outfit.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_outfit.autofit = False
        
        col_w = Inches(6.8 / len(outfits))
        for idx, card in enumerate(outfits):
            c = t_outfit.cell(0, idx)
            c.width = col_w
            set_cell_background(c, card.get("bg", "F8FAFC"))
            set_cell_margins(c, top=140, bottom=140, left=100, right=100)
            set_cell_border(c, top=dict(sz=6, color="CBD5E1"), bottom=dict(sz=6, color="CBD5E1"), left=dict(sz=6, color="CBD5E1"), right=dict(sz=6, color="CBD5E1"))
            
            cp = c.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(2)
            rn = cp.add_run(card.get("name", "成员"))
            rn.font.bold = True
            rn.font.size = Pt(10)
            
            cpt = c.add_paragraph()
            cpt.paragraph_format.space_before = Pt(0)
            cpt.paragraph_format.space_after = Pt(4)
            rt = cpt.add_run(card.get("tag", "穿搭角色"))
            rt.font.bold = True
            rt.font.size = Pt(8)
            rt.font.color.rgb = RGBColor(100, 116, 139)
            
            cpd = c.add_paragraph()
            cpd.paragraph_format.space_before = Pt(0)
            cpd.paragraph_format.space_after = Pt(0)
            rd = cpd.add_run(card.get("desc", ""))
            rd.font.size = Pt(8.5)
            rd.font.color.rgb = RGBColor(51, 65, 85)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Advance prep table
    prep_rows = data.get("prep_rows", [])
    if prep_rows:
        add_heading("行前关键节点与准备时间表", "⏰")
        t_prep = doc.add_table(rows=len(prep_rows) + 1, cols=4)
        t_prep.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_prep.autofit = False
        p_widths = [Inches(1.2), Inches(1.6), Inches(1.8), Inches(2.2)]
        
        headers = ["时间节点", "事项与目标", "平台 / 渠道", "核心实操要点"]
        for idx, h in enumerate(headers):
            c = t_prep.cell(0, idx)
            c.width = p_widths[idx]
            set_cell_background(c, "F1F5F9")
            set_cell_margins(c, top=100, bottom=100, left=80, right=80)
            set_cell_border(c, top=dict(sz=4, color="CBD5E1"), bottom=dict(sz=6, color="94A3B8"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(9)

        for r_idx, row in enumerate(prep_rows):
            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx in range(4):
                val = row[c_idx] if c_idx < len(row) else ""
                c = t_prep.cell(r_idx + 1, c_idx)
                c.width = p_widths[c_idx]
                set_cell_background(c, bg)
                set_cell_margins(c, top=80, bottom=80, left=80, right=80)
                set_cell_border(c, top=dict(sz=4, color="E2E8F0"), bottom=dict(sz=4, color="E2E8F0"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(val))
                r.font.size = Pt(8.5)
                if c_idx == 0: r.font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4. Day by day schedule cards
    days = data.get("days", [])
    if days:
        add_heading(f"{len(days)} 天逐日保姆级时间表", "📅")
        for day in days:
            t_day = doc.add_table(rows=1, cols=1)
            t_day.alignment = WD_TABLE_ALIGNMENT.CENTER
            t_day.autofit = False
            c_day = t_day.cell(0, 0)
            c_day.width = Inches(6.8)
            set_cell_background(c_day, "F8FAFC")
            set_cell_margins(c_day, top=140, bottom=140, left=160, right=160)
            set_cell_border(c_day, left=dict(sz=14, color=day.get("color_hex", "2563EB")), top=dict(sz=4, color="E2E8F0"), bottom=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
            
            p = c_day.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r_tag = p.add_run(f"【{day.get('day_tag', 'Day')}】")
            r_tag.font.bold = True
            r_tag.font.size = Pt(10.5)
            r_tag.font.color.rgb = RGBColor(15, 23, 42)
            
            r_sub = p.add_run(f" {day.get('day_title', '')}")
            r_sub.font.bold = True
            r_sub.font.size = Pt(9.5)
            r_sub.font.color.rgb = RGBColor(71, 85, 105)
            
            for t_time, t_desc in day.get("items", []):
                pi = c_day.add_paragraph()
                pi.paragraph_format.space_before = Pt(2)
                pi.paragraph_format.space_after = Pt(2)
                rt = pi.add_run(f"• {t_time}：")
                rt.font.bold = True
                rt.font.size = Pt(8.5)
                rt.font.color.rgb = RGBColor(30, 41, 59)
                
                rd = pi.add_run(str(t_desc))
                rd.font.size = Pt(8.5)
                rd.font.color.rgb = RGBColor(71, 85, 105)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # 5. Dining deals table
    deals = data.get("deals", [])
    if deals:
        add_heading("在地特色美食 ＆「老饕高分餐厅」指南", "🥢")
        t_deal = doc.add_table(rows=len(deals) + 1, cols=6)
        t_deal.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_deal.autofit = False
        d_widths = [Inches(1.3), Inches(1.6), Inches(0.7), Inches(0.8), Inches(0.6), Inches(1.8)]
        
        d_headers = ["餐厅名称与地点", "推荐特色硬菜", "参考消费", "优惠/预订", "推荐指数", "氛围与订座电话"]
        for idx, h in enumerate(d_headers):
            c = t_deal.cell(0, idx)
            c.width = d_widths[idx]
            set_cell_background(c, "F1F5F9")
            set_cell_margins(c, top=100, bottom=100, left=60, right=60)
            set_cell_border(c, top=dict(sz=4, color="CBD5E1"), bottom=dict(sz=6, color="94A3B8"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(8.5)

        for r_idx, row in enumerate(deals):
            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx in range(6):
                val = row[c_idx] if c_idx < len(row) else ""
                c = t_deal.cell(r_idx + 1, c_idx)
                c.width = d_widths[c_idx]
                set_cell_background(c, bg)
                set_cell_margins(c, top=60, bottom=60, left=60, right=60)
                set_cell_border(c, top=dict(sz=4, color="E2E8F0"), bottom=dict(sz=4, color="E2E8F0"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(val))
                r.font.size = Pt(8)
                if c_idx in (3, 4):
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(225, 29, 72)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 6. Parking (ONLY if parking data exists)
    parking = data.get("parking", [])
    if parking:
        add_heading("精准停车与自驾导航表", "🅿️")
        t_park = doc.add_table(rows=len(parking) + 1, cols=5)
        t_park.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_park.autofit = False
        p_w = [Inches(1.3), Inches(1.2), Inches(1.5), Inches(1.0), Inches(1.8)]
        p_hdrs = ["目的地", "推荐停车场", "导航关键字", "收费标准", "核心优势与设施"]
        for idx, h in enumerate(p_hdrs):
            c = t_park.cell(0, idx)
            c.width = p_w[idx]
            set_cell_background(c, "F1F5F9")
            set_cell_margins(c, top=100, bottom=100, left=60, right=60)
            set_cell_border(c, top=dict(sz=4, color="CBD5E1"), bottom=dict(sz=6, color="94A3B8"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(8.5)

        for r_idx, row in enumerate(parking):
            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx in range(5):
                val = row[c_idx] if c_idx < len(row) else ""
                c = t_park.cell(r_idx + 1, c_idx)
                c.width = p_w[c_idx]
                set_cell_background(c, bg)
                set_cell_margins(c, top=60, bottom=60, left=60, right=60)
                set_cell_border(c, top=dict(sz=4, color="E2E8F0"), bottom=dict(sz=4, color="E2E8F0"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(val))
                r.font.size = Pt(8)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 7. Souvenirs section (ONLY if souvenirs data exists)
    souvenirs = data.get("souvenirs", [])
    if souvenirs:
        add_heading("特色纪念品 ＆ 伴手礼淘货全攻略", "🎁")
        s_rows = max(1, (len(souvenirs) + 1) // 2)
        t_souv = doc.add_table(rows=s_rows, cols=2)
        t_souv.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_souv.autofit = False
        
        for idx, (s_name, s_desc) in enumerate(souvenirs):
            r_i = idx // 2
            c_i = idx % 2
            if r_i < s_rows:
                c = t_souv.cell(r_i, c_i)
                c.width = Inches(3.4)
                set_cell_background(c, "FFFBEB")
                set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                set_cell_border(c, top=dict(sz=4, color="FDE68A"), bottom=dict(sz=4, color="FDE68A"), left=dict(sz=4, color="FDE68A"), right=dict(sz=4, color="FDE68A"))
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                r_tag = p.add_run("【必淘】")
                r_tag.font.bold = True
                r_tag.font.size = Pt(8.5)
                r_tag.font.color.rgb = RGBColor(180, 83, 9)
                r_title = p.add_run(f" {s_name}")
                r_title.font.bold = True
                r_title.font.size = Pt(9.5)
                r_title.font.color.rgb = RGBColor(15, 23, 42)
                
                pd = c.add_paragraph()
                pd.paragraph_format.space_before = Pt(0)
                pd.paragraph_format.space_after = Pt(0)
                rd = pd.add_run(str(s_desc))
                rd.font.size = Pt(8.5)
                rd.font.color.rgb = RGBColor(71, 85, 105)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 8. Checklist
    checklist = data.get("checklist", [])
    if checklist:
        add_heading("行前打勾清单（已为您逐项核验备齐）", "🛒")
        chk_rows = max(1, (len(checklist) + 1) // 2)
        t_chk = doc.add_table(rows=chk_rows, cols=2)
        t_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_chk.autofit = False
        
        for idx, item in enumerate(checklist):
            r_i = idx // 2
            c_i = idx % 2
            if r_i < chk_rows:
                c = t_chk.cell(r_i, c_i)
                c.width = Inches(3.4)
                set_cell_background(c, "F8FAFC")
                set_cell_margins(c, top=60, bottom=60, left=100, right=100)
                set_cell_border(c, top=dict(sz=4, color="E2E8F0"), bottom=dict(sz=4, color="E2E8F0"), left=dict(sz=4, color="E2E8F0"), right=dict(sz=4, color="E2E8F0"))
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                r_box = p.add_run("☑ ")
                r_box.font.bold = True
                r_box.font.color.rgb = RGBColor(22, 163, 74)
                r_txt = p.add_run(str(item))
                r_txt.font.size = Pt(8.5)
                r_txt.font.color.rgb = RGBColor(51, 65, 85)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Ensure output parent dir exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated publication-grade Word docx at: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate publication-grade Word docx from travel markdown")
    parser.add_argument("--input", "-i", default="", help="Path to travel markdown file")
    parser.add_argument("--output", "-o", default="dist/travel_plan.docx", help="Path to output docx file")
    args = parser.parse_args()
    build_docx(args.input, args.output)
